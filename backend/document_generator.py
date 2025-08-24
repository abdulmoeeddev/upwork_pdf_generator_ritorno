from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from io import BytesIO
import json

def generate_word_document(proposal_json):
    """
    Generate Word document from proposal JSON
    """
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Upwork Proposal', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sections from JSON
        for section, content in proposal_json.items():
            if isinstance(content, dict):
                # Handle nested objects
                doc.add_heading(section.title().replace('_', ' '), level=1)
                for subsection, subcontent in content.items():
                    doc.add_heading(subsection.title().replace('_', ' '), level=2)
                    if isinstance(subcontent, dict):
                        for key, value in subcontent.items():
                            p = doc.add_paragraph()
                            p.add_run(f"{key.title().replace('_', ' ')}: ").bold = True
                            p.add_run(str(value))
                    else:
                        doc.add_paragraph(str(subcontent))
            else:
                doc.add_heading(section.title().replace('_', ' '), level=1)
                doc.add_paragraph(str(content))
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"Word generation error: {e}")
        # Return empty document on error
        doc = Document()
        doc.add_paragraph("Error generating document")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

def generate_pdf_document(proposal_json):
    """
    Generate PDF document from proposal JSON
    """
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center aligned
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12
        )
        
        content = []
        
        # Add title
        content.append(Paragraph('Upwork Proposal', title_style))
        content.append(Spacer(1, 20))
        
        # Add sections from JSON
        for section, section_content in proposal_json.items():
            content.append(Paragraph(section.title().replace('_', ' '), heading_style))
            
            if isinstance(section_content, dict):
                for subsection, subcontent in section_content.items():
                    content.append(Paragraph(subsection.title().replace('_', ' '), styles['Heading3']))
                    if isinstance(subcontent, dict):
                        for key, value in subcontent.items():
                            content.append(Paragraph(
                                f"<b>{key.title().replace('_', ' ')}:</b> {value}",
                                styles['BodyText']
                            ))
                    else:
                        content.append(Paragraph(str(subcontent), styles['BodyText']))
            else:
                content.append(Paragraph(str(section_content), styles['BodyText']))
            
            content.append(Spacer(1, 12))
        
        doc.build(content)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"PDF generation error: {e}")
        # Return error message PDF
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 700, "Error generating PDF document")
        c.save()
        buffer.seek(0)
        return buffer