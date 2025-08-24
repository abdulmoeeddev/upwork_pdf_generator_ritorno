from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from io import BytesIO
import json

class DocumentService:
    
    @staticmethod
    def generate_word_document(proposal_json, proposal_title="Upwork Proposal"):
        """
        Generate Word document from proposal JSON
        """
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add title
            title = doc.add_heading(proposal_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add a line break
            doc.add_paragraph()
            
            # Add sections from JSON
            for section_key, content in proposal_json.items():
                if section_key == 'revision_notes':
                    continue  # Skip revision notes in final document
                    
                section_title = DocumentService._format_section_title(section_key)
                doc.add_heading(section_title, level=1)
                
                if isinstance(content, dict):
                    DocumentService._add_dict_content_to_doc(doc, content)
                elif isinstance(content, list):
                    DocumentService._add_list_content_to_doc(doc, content)
                else:
                    doc.add_paragraph(str(content))
                
                # Add spacing between sections
                doc.add_paragraph()
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            print(f"Word generation error: {e}")
            return DocumentService._generate_error_word_doc(str(e))
    
    @staticmethod
    def generate_pdf_document(proposal_json, proposal_title="Upwork Proposal"):
        """
        Generate PDF document from proposal JSON
        """
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=30,
                alignment=1,  # Center aligned
                textColor=HexColor('#1f4788')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=20,
                textColor=HexColor('#1f4788')
            )
            
            subheading_style = ParagraphStyle(
                'CustomSubheading',
                parent=styles['Heading3'],
                fontSize=12,
                spaceAfter=8,
                spaceBefore=12,
                textColor=HexColor('#2c5aa0')
            )
            
            content = []
            
            # Add title
            content.append(Paragraph(proposal_title, title_style))
            content.append(Spacer(1, 20))
            
            # Add sections from JSON
            for section_key, section_content in proposal_json.items():
                if section_key == 'revision_notes':
                    continue  # Skip revision notes
                    
                section_title = DocumentService._format_section_title(section_key)
                content.append(Paragraph(section_title, heading_style))
                
                if isinstance(section_content, dict):
                    DocumentService._add_dict_content_to_pdf(content, section_content, styles, subheading_style)
                elif isinstance(section_content, list):
                    DocumentService._add_list_content_to_pdf(content, section_content, styles)
                else:
                    content.append(Paragraph(str(section_content), styles['BodyText']))
                
                content.append(Spacer(1, 15))
            
            doc.build(content)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            print(f"PDF generation error: {e}")
            return DocumentService._generate_error_pdf_doc(str(e))
    
    @staticmethod
    def _format_section_title(section_key):
        """Format section key to readable title"""
        return section_key.replace('_', ' ').title()
    
    @staticmethod
    def _add_dict_content_to_doc(doc, content_dict):
        """Add dictionary content to Word document"""
        for key, value in content_dict.items():
            if isinstance(value, dict):
                doc.add_heading(DocumentService._format_section_title(key), level=2)
                DocumentService._add_dict_content_to_doc(doc, value)
            elif isinstance(value, list):
                doc.add_heading(DocumentService._format_section_title(key), level=2)
                DocumentService._add_list_content_to_doc(doc, value)
            else:
                p = doc.add_paragraph()
                p.add_run(f"{DocumentService._format_section_title(key)}: ").bold = True
                p.add_run(str(value))
    
    @staticmethod
    def _add_list_content_to_doc(doc, content_list):
        """Add list content to Word document"""
        for item in content_list:
            p = doc.add_paragraph(str(item), style='List Bullet')
    
    @staticmethod
    def _add_dict_content_to_pdf(content, content_dict, styles, subheading_style):
        """Add dictionary content to PDF"""
        for key, value in content_dict.items():
            if isinstance(value, dict):
                content.append(Paragraph(DocumentService._format_section_title(key), subheading_style))
                DocumentService._add_dict_content_to_pdf(content, value, styles, subheading_style)
            elif isinstance(value, list):
                content.append(Paragraph(DocumentService._format_section_title(key), subheading_style))
                DocumentService._add_list_content_to_pdf(content, value, styles)
            else:
                content.append(Paragraph(
                    f"<b>{DocumentService._format_section_title(key)}:</b> {value}",
                    styles['BodyText']
                ))
    
    @staticmethod
    def _add_list_content_to_pdf(content, content_list, styles):
        """Add list content to PDF"""
        for item in content_list:
            content.append(Paragraph(f"• {item}", styles['BodyText']))
    
    @staticmethod
    def _generate_error_word_doc(error_message):
        """Generate error Word document"""
        doc = Document()
        doc.add_heading('Document Generation Error', 0)
        doc.add_paragraph(f"An error occurred while generating the document: {error_message}")
        doc.add_paragraph("Please contact support if this issue persists.")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def _generate_error_pdf_doc(error_message):
        """Generate error PDF document"""
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 700, "Document Generation Error")
        c.drawString(100, 670, f"Error: {error_message}")
        c.drawString(100, 640, "Please contact support if this issue persists.")
        c.save()
        buffer.seek(0)
        return buffer